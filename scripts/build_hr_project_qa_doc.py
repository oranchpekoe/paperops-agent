from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/hr_project_angle_qa.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 32, 44)
MUTED = RGBColor(90, 98, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WARN = "FFF4D6"


def set_run_font(run, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for k, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{k}"))
        if node is None:
            node = OxmlElement(f"w:{k}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def borders(table, color="D7DCE2", size="6"):
    tbl_pr = table._tbl.tblPr
    tb = tbl_pr.first_child_found_in("w:tblBorders")
    if tb is None:
        tb = OxmlElement("w:tblBorders")
        tbl_pr.append(tb)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tb.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tb.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def title(doc):
    p = doc.add_paragraph()
    r = p.add_run("从项目角度准备 HR 问题与答案")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 37, 69)
    set_run_font(r)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run("围绕 Multi-Mode Agent Framework 项目，准备 HR 初筛、项目真实性、学历弱势、岗位匹配和压力追问")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    set_run_font(r)


def callout(doc, head, body, fill=LIGHT_GRAY):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    borders(t, size="4")
    table_width(t, [6.5])
    c = t.cell(0, 0)
    shade(c, fill)
    p = c.paragraphs[0]
    r = p.add_run(head)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    set_run_font(r)
    p.add_run("\n" + body)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def qa(doc, q, answer, point=None, avoid=None):
    doc.add_paragraph(q, style="Heading 3")
    p = doc.add_paragraph()
    p.add_run("建议回答：").bold = True
    p.add_run(answer)
    if point:
        p = doc.add_paragraph()
        p.add_run("回答重点：").bold = True
        p.add_run(point)
    if avoid:
        p = doc.add_paragraph()
        p.add_run("不要这样说：").bold = True
        p.add_run(avoid)


def matrix(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    borders(t)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        shade(cell, LIGHT_BLUE)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    table_width(t, widths)


def build():
    doc = Document()
    style_doc(doc)
    title(doc)

    callout(
        doc,
        "这份文档的用法",
        "HR 不是技术面试官，但会判断三件事：你是不是靠谱、项目是不是你真的做过、你和岗位是不是匹配。回答时少堆术语，多讲动机、贡献、取舍、复盘和稳定性。技术词只作为证据，不作为炫耀。",
        WARN,
    )

    doc.add_heading("1. 你的项目在 HR 面前应该怎么定位", level=1)
    matrix(
        doc,
        ["HR 关心点", "你要传递的信号", "项目证据", "一句话表达"],
        [
            ("真实性", "不是简历包装", "能讲清 clone 基础和自己改造部分", "我基于开源项目做了系统化二次改造。"),
            ("学习能力", "能快速补新技术", "LangGraph、MCP、RAG、Memory、Eval", "我通过项目把 Agent 工程链路串起来了。"),
            ("工程意识", "不只调 demo", "测试、benchmark、streaming、bug 修复", "我更关注可测试、可复盘、可扩展。"),
            ("诚实边界", "不硬吹", "记录 human-in-loop、sandbox、日志短板", "我知道它还不是生产级，也知道下一步怎么补。"),
            ("岗位匹配", "能做实习基础活", "工具接入、评测、文档、排查", "我愿意从真实业务里的基础工程任务做起。"),
        ],
        [1.15, 1.45, 1.9, 2.0],
    )

    doc.add_heading("2. HR 初筛高频问题", level=1)
    qa(
        doc,
        "Q1：请你简单介绍一下自己。",
        "您好，我是双非硕士背景，最近主要在准备 Agent 应用工程方向的实习。我做了一个基于 LangGraph 的 Multi-Mode Agent Framework 项目，重点不是从零训练模型，而是围绕 Agent 落地链路做工程化改造：包括多模式路由、MCP 工具接入、RAG 文档检索、三层记忆、Supervisor 多 Agent 协作、streaming 和 benchmark 评测。我的定位更偏应用工程，想在实习里继续把工具接入、评测、日志、稳定性这些能力补扎实。",
        "30 秒内交代背景、方向、项目主线和岗位定位。",
    )
    qa(
        doc,
        "Q2：为什么想找 Agent 实习？",
        "我对 Agent 感兴趣不是因为概念热，而是因为它把模型能力和工程落地结合得很紧。一个 Agent 要真的能用，除了模型本身，还要处理工具调用、知识检索、记忆、上下文、评测和安全。我这个项目就是按这个链路去补的，所以我希望进入真实业务团队，看看这些能力在企业场景里怎么落地。",
        "强调“落地链路”，不要只说“AI 很火”。",
    )
    qa(
        doc,
        "Q3：你为什么不投普通后端，而是投 Agent？",
        "我并不排斥后端，实际上 Agent 应用工程里也有很多后端能力，比如接口封装、异步调用、状态管理、日志、测试和工具服务接入。区别是 Agent 多了模型行为、prompt、RAG 和评测这些问题。我现在的项目正好把后端工程和 LLM 应用结合起来，所以我更想先尝试 Agent 方向。",
    )
    qa(
        doc,
        "Q4：你能实习多久？到岗时间怎么样？",
        "建议你按自己的真实情况回答。模板：如果团队匹配，我可以尽快开始实习，预计可以持续 X 个月，每周 X 天。对我来说，第一份 Agent 实习比短期项目更重要，所以我希望能稳定参与一段完整的业务迭代。",
        "稳定性是 HR 很看重的，不要含糊。",
    )
    qa(
        doc,
        "Q5：你对实习岗位有什么期待？",
        "我希望能参与真实 Agent 应用的落地工作，比如工具接入、RAG 流程、评测集建设、prompt 调优、日志排查和小模块开发。我不期待一进去就做很核心的架构设计，更希望先把团队需要的基础任务做好，在真实项目里提升。",
        "这句话很适合中厂：谦逊、可用、稳定。",
    )

    doc.add_heading("3. 项目真实性与贡献度追问", level=1)
    qa(
        doc,
        "Q1：这个项目是你自己做的吗？",
        "我会诚实区分：项目基础是从 GitHub clone 下来的 LangGraph ReAct 模板，但后续我做了比较系统的改造。我的贡献主要包括四类：第一，把单一 ReAct 扩展成 ReAct、Reflection、Plan-Solve、Supervisor 四种模式；第二，接入 MCP、RAG 和记忆系统；第三，补 streaming、benchmark 和测试；第四，记录和修复了一些工程问题，比如 structured output、benchmark 记忆污染和 ASGI blocking。它不是完全从零原创，但改造过程是我自己理解和推进的。",
        "诚实承认 clone，马上转向具体改造。",
        "不要说“全部是我原创的”。这很容易被追问穿。",
    )
    qa(
        doc,
        "Q2：你在项目里最核心的贡献是什么？",
        "我最核心的贡献是把它从一个偏 demo 的 Agent 项目，往工程化展示项目推进。比如我不是只让模型能回答问题，而是补了路由、工具、RAG、记忆、评测和流式输出这些完整链路。尤其是 benchmark 和记忆隔离，让项目不只是能跑，而是能被验证和复盘。",
        "HR 不一定懂技术，但能听懂“从 demo 到可验证”。",
    )
    qa(
        doc,
        "Q3：如果面试官说这只是 GitHub 项目包装，你怎么回应？",
        "我会说：这个质疑是合理的，所以我不会把它包装成完全原创框架。我想展示的是二次工程化能力：能读懂开源项目，能识别短板，能补功能、补测试、补文档，能讲清楚每次改造的原因和效果。实习阶段我觉得这种能力也很重要，因为真实工作里很多任务就是在已有系统上迭代。",
        "把“不是原创”转成“会在已有系统上迭代”。",
    )
    qa(
        doc,
        "Q4：项目里最能说明你能力的地方是什么？",
        "我会选三个点讲：第一是多模式路由，说明我理解不同 Agent 架构适合不同任务；第二是记忆和 benchmark 隔离，说明我考虑到长期运行中的污染问题；第三是评测框架，说明我不是凭感觉判断效果，而是尝试用固定 case 做回归验证。",
    )
    qa(
        doc,
        "Q5：项目目前还有哪些不足？",
        "不足我会主动承认：它还不是生产级系统。比如 human-in-the-loop 还没做，python_repl 的 sandbox 只适合 demo，不适合公开环境；结构化日志和 metrics 还不完善；RAG 也还没有混合检索和 rerank。我已经把这些记录在 gaps 文档里，后续会按安全、可观测性、检索质量的优先级继续补。",
        "主动承认短板，会比硬吹更可信。",
    )

    doc.add_heading("4. 学历弱势与竞争力问题", level=1)
    qa(
        doc,
        "Q1：你是双非硕士，和名校同学比优势在哪里？",
        "我承认学校背景不是我的优势，所以我会尽量用具体项目和复盘来证明自己。我这个项目不是简单跑通，而是围绕 Agent 工程链路做了改造，也记录了问题和不足。我觉得自己的优势是自驱、能持续补知识、愿意做工程细节，也愿意从基础任务开始稳定交付。",
        "承认弱势，不卖惨；转向可验证行为。",
    )
    qa(
        doc,
        "Q2：你觉得自己竞争力弱吗？",
        "如果只看学历，我确实不是最强的一档。但我不想只用学历定义自己。我现在能拿出来的是一个比较完整的 Agent 工程项目、清晰的复盘文档，以及愿意继续补基础的态度。对实习岗位来说，我希望证明自己能学、能改、能测、能配合团队推进任务。",
    )
    qa(
        doc,
        "Q3：你为什么觉得公司应该给你机会？",
        "因为我对自己的定位比较清楚：我不是来证明自己已经是专家，而是希望作为一个能快速学习、认真交付的实习生加入团队。我有 Agent 项目基础，也愿意做工具接入、评测、数据整理、文档和排查这些具体工作。如果团队需要一个踏实推进应用工程的实习生，我觉得我能匹配。",
        "这类回答要稳，不要过度自夸。",
    )
    qa(
        doc,
        "Q4：你的基础是不是不够扎实？",
        "有些底层方向我还在补，比如更系统的 RAG 评测、Agent 训练、生产级安全和可观测性。但我不会回避这些短板。我现在的做法是通过项目把问题暴露出来，再逐个补概念和实践。相比说自己什么都会，我更希望展示持续学习和复盘能力。",
        "这正好回应你现在的状态，真实但不虚。",
    )

    doc.add_heading("5. 项目价值与岗位匹配", level=1)
    qa(
        doc,
        "Q1：这个项目和我们公司的 Agent 岗位有什么关系？",
        "我理解真实公司的 Agent 岗位不只是写 prompt，而是要把模型接入业务流程。我的项目虽然是个人项目，但覆盖了很多相似环节：任务路由、工具调用、RAG、记忆、流式输出、评测和问题复盘。所以它能证明我对 Agent 应用工程链路有整体理解，入职后能更快接上团队的工具和业务。",
    )
    qa(
        doc,
        "Q2：如果实习中让你做评测集、文档、工具接入，你愿意吗？",
        "愿意。我知道 Agent 落地里这些工作非常重要。没有评测就不知道改动是否有效，没有文档团队就难以协作，没有工具接入 Agent 就不能真正执行任务。我不排斥这些基础工作，反而希望通过它们理解真实系统。",
    )
    qa(
        doc,
        "Q3：你更想做研究还是工程？",
        "我目前更偏工程落地。研究方向我会关注，但我的优势和兴趣是在已有模型基础上做应用：接工具、做 RAG、设计流程、补评测、优化体验和稳定性。所以我投的是 Agent 应用工程实习，而不是大模型算法研究岗。",
    )
    qa(
        doc,
        "Q4：你希望 mentor 怎么带你？",
        "我希望 mentor 能在方向和标准上给我反馈，比如告诉我任务的业务目标、代码规范和验收标准。我自己会主动拆任务、查文档、做记录，遇到卡点先自己定位，再带着现象和尝试过的方法去请教。",
        "HR 喜欢听到“先自查，再请教”。",
    )
    qa(
        doc,
        "Q5：你入职后第一个月能做什么？",
        "我觉得可以从三类事情开始：第一，熟悉团队现有 Agent 链路和工具；第二，参与一些低风险模块，比如工具封装、prompt case、评测脚本、文档整理；第三，跟着已有 issue 做 bug 定位和小功能迭代。我的目标是先稳定交付，再逐步承担更完整的模块。",
    )

    doc.add_heading("6. HR 压力追问：更尖锐的问题", level=1)
    qa(
        doc,
        "Q1：你项目这么多功能，会不会只是堆概念？",
        "这个风险确实存在，所以我面试时不会只列技术名词。我会重点讲每个功能解决的问题：路由解决不同任务走不同策略，RAG 解决外部知识，记忆解决跨会话上下文，benchmark 解决效果验证，MCP 解决工具扩展。它们不是为了堆概念，而是围绕 Agent 落地链路补齐能力。",
    )
    qa(
        doc,
        "Q2：你没有真实上线经验，怎么证明能适应公司项目？",
        "我没有把个人项目说成生产项目，这点我会承认。但我通过项目至少接触了生产化会遇到的方向，比如安全沙箱、human-in-the-loop、日志、评测、成本和延迟。我还没有完整实践，但知道问题在哪里，也愿意在团队规范下学习和补齐。",
    )
    qa(
        doc,
        "Q3：你是不是只会用 AI 帮你写代码？",
        "我会使用 AI 工具提升效率，但我不会把它当成替代理解的东西。这个项目里我需要理解 LangGraph 的状态流转、工具调用、记忆污染、benchmark 评估这些问题，否则无法解释改造原因和取舍。AI 可以帮我写样板代码，但项目结构、问题定位和面试复盘必须自己理解。",
        "这个问题现在很常见，回答要大方，不要防御。",
    )
    qa(
        doc,
        "Q4：如果技术面问得很深你答不上来怎么办？",
        "我会先承认这个点我没有深入实践，然后说出我已有的理解和下一步学习路径。比如被问到 Agent 训练，我会说明我的项目没有训练模型，主要是应用工程；如果要继续做，会先收集高质量轨迹和评测指标，再考虑 fine-tune 或偏好优化。我不会硬编。",
    )
    qa(
        doc,
        "Q5：你是不是对岗位理解还不够成熟？",
        "我确实还在学习，但我对 Agent 实习的理解已经不只停留在 prompt。我知道真实岗位可能会涉及工具接入、RAG、评测、日志、数据处理、接口联调和业务落地。我愿意从这些具体任务做起，而不是只期待做很抽象的模型工作。",
    )

    doc.add_heading("7. HR 反问与收尾话术", level=1)
    bullets(
        doc,
        [
            "这个岗位更偏 Agent 应用工程，还是偏平台基础设施？",
            "实习生入职后，前 1 个月通常会参与什么类型的任务？",
            "团队目前 Agent 项目里最关注的是 RAG 效果、工具调用，还是评测和稳定性？",
            "团队会有 mentor 带实习生做 code review 和任务拆解吗？",
            "如果我想在入职前补准备，您建议我重点补哪块？",
        ]
    )
    callout(
        doc,
        "面试结束收尾",
        "今天交流后，我感觉这个岗位和我想做的 Agent 应用工程比较匹配。我现在还不是很成熟的工程师，但我有项目基础，也愿意做工具接入、评测、文档、排查这些具体工作。如果有机会进入团队，我会先把基础任务稳定交付，再逐步承担更复杂的模块。",
        WARN,
    )

    doc.add_heading("8. 最后速背：项目角度 HR 答案骨架", level=1)
    bullets(
        doc,
        [
            "项目不是从零原创，但我做了系统化二次改造。",
            "我的核心价值是从 demo 往工程化推进：路由、工具、RAG、记忆、评测、streaming。",
            "学历不是优势，所以我用项目、复盘、稳定性和学习能力补足。",
            "我不把自己包装成专家，定位是能快速学习和交付的 Agent 应用工程实习生。",
            "真实工作里基础任务很重要，我愿意做工具接入、评测、文档、排查和接口联调。",
            "项目还有不足：human-in-the-loop、sandbox、结构化日志、混合检索、生产权限都需要补。",
        ]
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("HR 项目角度问答 | Multi-Mode Agent Framework | Agent 实习准备")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    set_run_font(r)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
