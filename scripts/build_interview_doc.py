from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/agent_internship_interview_qa.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 32, 44)
MUTED = RGBColor(90, 98, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
    table.autofit = False


def set_table_borders(table, color="D7DCE2", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Agent 实习面试问题与参考答案")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)
    set_run_font(run)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("适用于双非硕士候选人：用真实项目改造、工程意识和复盘能力争取中厂 Agent 实习")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED
    set_run_font(run)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, [1.45, 5.05])
    rows = [
        ("使用方法", "先背第 1、2、3 部分，再按第 4、5 部分做追问演练。不要逐字背答案，要背逻辑。"),
        ("个人定位", "双非硕士不回避，重点突出自驱学习、工程落地、能把开源项目改造成可讲述系统。"),
        ("项目主线", "不是“我调了个 Agent demo”，而是“我围绕多模式调度、工具生态、记忆、评估和可观测性做工程化改造”。"),
    ]
    for row, (k, v) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(v)


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, color="D7DCE2", size="4")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run("\n" + body)


def add_qa(doc, question, answer, follow=None, danger=None):
    p = doc.add_paragraph(style="Heading 3")
    p.add_run(question)
    p = doc.add_paragraph()
    p.add_run("参考回答：").bold = True
    p.add_run(answer)
    if follow:
        p = doc.add_paragraph()
        p.add_run("可能追问：").bold = True
        p.add_run(follow)
    if danger:
        p = doc.add_paragraph()
        p.add_run("注意避坑：").bold = True
        p.add_run(danger)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. 面试总策略：双非硕士怎么讲才不弱", level=1)
    add_callout(
        doc,
        "核心原则",
        "不要把自己包装成“天才型选手”，而是讲成“能自驱补齐能力、能把开源工程改造成可验证系统、愿意做脏活累活的工程型实习生”。中厂实习更看重你能不能快速上手、稳定交付、持续学习。",
    )
    doc.add_heading("推荐定位", level=2)
    add_bullets(
        doc,
        [
            "学历上我不是最强的一档，所以我会主动用项目和复盘证明学习能力。",
            "这个项目不是从零发明框架，而是基于开源模板做系统化改造，重点体现工程理解和落地意识。",
            "我能接受实习中做工具接入、评测、prompt 调优、问题定位、文档整理这类基础但重要的工作。",
            "我比较适合 Agent 应用工程岗位：懂一点模型行为，也愿意把框架、工具、测试、日志这些工程细节补起来。",
        ]
    )
    doc.add_heading("不要这样说", level=2)
    add_bullets(
        doc,
        [
            "不要说：我学历比较差，所以只能靠项目弥补。",
            "不要说：这个项目基本是 GitHub 上的，我只是改了一点。",
            "不要说：我什么 Agent 技术都会，但还没上线过。",
            "不要说：我只想做核心算法，不太想做工程杂活。",
        ]
    )
    doc.add_heading("推荐这样说", level=2)
    add_callout(
        doc,
        "30 秒版本",
        "我本身是双非硕士，学校背景不算强，所以我准备实习时没有只停留在刷概念，而是选了一个 LangGraph Agent 项目做系统化改造。我主要补了多模式路由、MCP 工具扩展、RAG、三层记忆、Supervisor 多 Agent 协作、streaming 和 benchmark。这个项目还不是生产级系统，但它能体现我对 Agent 工程链路的理解，也能说明我愿意把一个 demo 往可测试、可复盘、可扩展的方向推进。",
    )

    doc.add_heading("2. 两分钟项目介绍", level=1)
    add_callout(
        doc,
        "可直接背诵",
        "我做的是一个基于 LangGraph 的多模式 Agent 框架。它不是只有一个 ReAct，而是在入口用 LLM Router 判断用户意图，然后路由到 ReAct、Reflection、Plan-Solve 或 Supervisor-Worker 四种模式。ReAct 处理简单问答和工具调用，Reflection 处理写作分析和自我修正，Plan-Solve 处理多步骤任务，Supervisor-Worker 处理搜索、分析、计算混合的复杂任务。\n\n工程上我主要做了几块：第一是 MCP 协议接入，让外部工具可以动态加载；第二是 RAG，本地文档通过 Chroma 做向量检索；第三是三层记忆，包含短期会话、长期 Chroma 记忆和摘要压缩；第四是 Supervisor 多 Agent 协作，用 structured output 约束决策；第五是 streaming 和 benchmark，用来验证路由、工具调用、回答质量和多步推理。\n\n我最想强调的是，这个项目不是为了展示一个炫技 demo，而是为了练习 Agent 工程化：怎么路由、怎么接工具、怎么做记忆、怎么防止评测污染、怎么写测试和 benchmark、怎么承认并记录生产化短板。",
    )

    doc.add_heading("3. HR 初筛问题 + 答案", level=1)
    add_qa(
        doc,
        "Q1：请简单介绍一下自己，以及为什么想找 Agent 方向实习？",
        "我是双非硕士背景，方向上更偏工程应用。最近重点学习 Agent 工程，做了一个基于 LangGraph 的多模式 Agent 项目。选择 Agent 方向是因为它既需要理解模型能力，也需要后端工程、工具接入、RAG、评测和可观测性这些落地能力。我觉得自己目前最适合从 Agent 应用工程实习切入，在真实业务里把这些能力继续补扎实。",
        "为什么不是算法岗？可以回答：我更偏应用落地和工程实现，不把自己包装成训练大模型的人。",
    )
    add_qa(
        doc,
        "Q2：你的学历不是特别有竞争力，你怎么弥补？",
        "我承认学历不是优势，所以我会用更具体的项目和复盘来弥补。比如这个项目我不是只跑通 demo，而是围绕路由、工具、记忆、评测、streaming 做了完整改造，并记录了问题、修复和剩余短板。我希望面试官看到的是我的学习速度、工程意识和持续推进能力。",
        danger="不要贬低学校，也不要过度解释高考、考研经历。重点放在“现在能交付什么”。",
    )
    add_qa(
        doc,
        "Q3：你这个项目是 clone 的，怎么证明是你自己的能力？",
        "我会明确区分：基础模板来自开源项目，但我的价值在于二次工程化改造。比如我补了 MCP demo server、RAG 持久化、三层记忆、Supervisor structured output、benchmark 隔离、streaming、单元测试和优化文档。面试时我可以具体讲每个模块改了什么、为什么改、遇到什么 bug、怎么验证。这比笼统说“我从零写了一个框架”更真实。",
    )
    add_qa(
        doc,
        "Q4：你相比其他候选人有什么优势？",
        "我的优势不是学历，而是我愿意把一个开源 demo 往工程系统推进。我不会只讲概念，会讲清楚路由、工具调用、记忆污染、structured output、评测、blocking 这些具体问题。实习里很多工作其实需要耐心接工具、跑评测、定位 bug、写文档，我对这些事情有准备，也愿意做。",
    )
    add_qa(
        doc,
        "Q5：你能接受实习中做比较基础的任务吗？",
        "可以。我理解 Agent 落地不只是写复杂 prompt，还包括工具封装、接口联调、数据清洗、评测集维护、日志排查和文档沉淀。我希望先在真实项目里把这些基础能力做好，再逐步参与更复杂的 Agent 架构设计。",
    )

    doc.add_heading("4. 技术面核心问题 + 答案", level=1)
    qa_items = [
        (
            "Q1：为什么设计四种 Agent 模式，而不是一个通用 prompt？",
            "通用 prompt 的问题是边界不清、token 成本高、调试困难。四种模式拆开后，每种模式有独立控制流：ReAct 负责工具增强问答，Reflection 负责自我批判和改写，Plan-Solve 负责多步骤规划，Supervisor 负责多专家协作。这样路由、日志、评测和优化都更清晰。",
            "如果 Router 错了怎么办？回答：靠 benchmark 发现错误样例，再优化 prompt 或加入 fallback。",
        ),
        (
            "Q2：Router 是怎么工作的？",
            "入口处会把用户 query 交给 LLM Router，让它输出 react、reflection、plan_solve 或 supervisor。然后 LangGraph 根据 mode 字段走条件边进入对应子图。它不是简单 regex，但我也保留了一些解析和兜底逻辑，避免模型输出格式不稳定导致流程崩掉。",
            None,
        ),
        (
            "Q3：Supervisor-Worker 的价值是什么？",
            "它适合混合型任务，比如既要查资料又要计算再分析。Supervisor 不直接干所有事，而是决定下一步委派给 Researcher、Analyst 或 Executor。这样可以把搜索、推理、计算职责拆开，也更容易限制每个专家的工具权限和迭代轮数。",
            None,
        ),
        (
            "Q4：为什么 Supervisor 决策要用 structured output？",
            "早期只靠文本解析，模型可能输出不规范文本，导致解析失败或默认结束。structured output 用 Pydantic schema 约束 action、reason 等字段，相当于从协议层要求模型输出可解析结构。对不支持的模型，我保留普通文本解析作为 fallback。",
            None,
        ),
        (
            "Q5：RAG 是怎么实现的？",
            "文档放在 docs 目录，支持 txt、md、pdf。加载后用 RecursiveCharacterTextSplitter 切块，再用 OpenAI-compatible embedding 写入 Chroma。retrieve 工具被 Agent 调用时会懒加载索引，后续复用同一个向量库，避免每次请求重建。",
            None,
        ),
        (
            "Q6：为什么选择 Chroma，而不是 FAISS？",
            "主要是 Windows 兼容性和持久化。FAISS 在 Windows 上经常涉及编译和环境问题，Chroma 纯 Python 体验更好，也原生支持 persist_directory。项目里还复用了 Chroma 做长期记忆，用不同 collection 隔离 RAG 文档和 agent_memory。",
            None,
        ),
        (
            "Q7：三层记忆分别解决什么问题？",
            "短期记忆解决当前会话上下文，长期记忆解决跨会话偏好和事实召回，摘要记忆解决上下文过长的问题。三层不是重复，而是对应不同时间尺度：当前对话、跨会话、超长上下文压缩。",
            None,
        ),
        (
            "Q8：benchmark 为什么会污染记忆？",
            "benchmark 里的合成问题如果被长期记忆系统存进去，后续真实对话可能召回这些假偏好或假事实，影响评测和用户体验。所以我加了 benchmark_mode 标志、ContextVar 和兜底模式匹配，确保 eval 不写入长期记忆。",
            None,
        ),
        (
            "Q9：MCP 接入解决了什么问题？",
            "MCP 让工具扩展不再都写死在 tools.py 里。外部 MCP server 可以动态提供工具，Agent 通过统一 get_all_tools 拿到内置工具和外部工具。单个 MCP server 失败时返回空工具或跳过，不影响主流程启动。",
            None,
        ),
        (
            "Q10：这个项目最大的问题是什么？",
            "它还不是生产级。主要短板是 human-in-the-loop 缺失、python_repl 的 eval sandbox 不够安全、结构化日志和 metrics 不完善、RAG 还不是混合检索、Plan-Solve 没有并行执行。但我把这些短板记录在 gaps 文档里，并能说明优先级和改造方案。",
            None,
        ),
    ]
    for q, a, f in qa_items:
        add_qa(doc, q, a, f)

    doc.add_heading("5. 项目真实性与改造价值：高压追问", level=1)
    add_qa(
        doc,
        "Q1：你是不是只是把 README 包装了一下？",
        "不是。我会用具体改造点回答：比如 supervisor 决策从文本解析迁移到 structured output，mini ReAct loop 抽成公共函数，benchmark_mode 防止评测污染长期记忆，stream.py 提供 token 流式输出，MCP demo server 验证外部工具接入。每个点都能对应代码文件、问题背景和验证方式。",
    )
    add_qa(
        doc,
        "Q2：如果面试官说这些功能网上都有，你的价值在哪里？",
        "我不会说这些概念是我发明的。我的价值是把这些概念组合成一个可运行、可测试、可复盘的工程项目，并且能讲清楚取舍。实习生阶段，能理解成熟方案、改造开源项目、定位问题、补测试和文档，本身就是团队需要的能力。",
    )
    add_qa(
        doc,
        "Q3：你怎么证明你不是只会调包？",
        "我会从 bug 和取舍讲，而不是只列技术栈。比如 structured output 解决解析脆弱，benchmark 隔离解决记忆污染，asyncio.to_thread 解决 ASGI blocking，Chroma collection 隔离复用基础设施。这些问题不是调包能自动解决的，需要理解框架运行时和工程边界。",
    )

    doc.add_heading("6. 可以主动讲的 6 个亮点故事", level=1)
    stories = [
        ("structured output", "Supervisor 文本解析不稳，迁移到 Pydantic structured output，同时保留 fallback。体现协议约束意识。"),
        ("ASGI blocking", "LangGraph dev 检测到同步阻塞，定位到 Chroma/tiktoken 链路，用 asyncio.to_thread 处理。体现问题定位能力。"),
        ("三层记忆", "不是简单把历史塞进 prompt，而是短期、长期、摘要分层，并处理 benchmark 污染。体现系统设计意识。"),
        ("MCP 工具生态", "工具不只写死在项目内，而是支持外部 server 动态接入，失败时优雅降级。体现扩展性意识。"),
        ("benchmark", "不只看 demo 成功，而是做 routing、tool use、quality、multi-step、memory 多维评估。体现质量意识。"),
        ("诚实记录短板", "Human-in-the-loop、eval sandbox、日志 metrics、混合检索都没有硬吹，而是记录方案。体现工程判断。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, [1.7, 4.8])
    table.rows[0].cells[0].text = "亮点"
    table.rows[0].cells[1].text = "怎么讲"
    for c in table.rows[0].cells:
        set_cell_shading(c, LIGHT_BLUE)
        c.paragraphs[0].runs[0].bold = True
    for k, v in stories:
        row = table.add_row()
        row.cells[0].text = k
        row.cells[1].text = v
    set_table_width(table, [1.7, 4.8])

    doc.add_heading("7. 反问 HR / 面试官", level=1)
    add_numbered(
        doc,
        [
            "这个岗位更偏 Agent 应用开发，还是偏平台/框架建设？",
            "实习生进去后，主要会做工具接入、评测、prompt 调优，还是后端工程开发？",
            "团队现在使用 LangGraph、LangChain、Dify，还是自研 Agent 框架？",
            "团队有没有自己的 Agent 评测体系？会不会关注工具调用成功率、成本和延迟？",
            "这个岗位会接触 RAG、function calling、MCP、多 Agent 协作这些方向吗？",
            "对实习生最看重的是工程能力、模型理解，还是业务落地能力？",
            "如果我入职，前 1 个月通常会负责什么类型的任务？",
            "您觉得我这个项目和岗位要求相比，最需要补强的地方是什么？",
        ]
    )

    doc.add_heading("8. 面试前一天背诵清单", level=1)
    add_bullets(
        doc,
        [
            "背熟两分钟项目介绍，但不要像背稿，要能根据面试官打断继续讲。",
            "准备 3 个具体代码文件：graph.py、tools.py、memory.py 或 supervisor.py。",
            "准备 2 个 bug 故事：structured output、ASGI blocking、benchmark 污染任选两个。",
            "准备 1 个诚实短板：eval sandbox 或 human-in-the-loop。",
            "准备 1 个学历回答：承认弱势，转向项目、复盘、自驱和可交付。",
            "准备 3 个反问：岗位方向、技术栈、实习任务。",
        ]
    )

    doc.add_heading("9. 最后给自己的话术底线", level=1)
    add_callout(
        doc,
        "底线话术",
        "我不是名校背景，也不把这个项目包装成完全原创。但我能清楚说出自己基于开源项目做了哪些工程化改造、为什么要做、遇到什么问题、怎么验证，以及下一步怎么补生产化短板。对实习岗位来说，我希望证明的是：我能学、能改、能测、能复盘，也愿意从真实业务里的基础工作做起。",
        fill="FFF4D6",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Agent 实习面试准备 | 问题 + 答案 | 双非硕士项目表达策略")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
