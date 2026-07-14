from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/agent_foundation_interview_cheatsheet.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 32, 44)
MUTED = RGBColor(90, 98, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WARN = "FFF4D6"


def set_run_font(run, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color="D7DCE2", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    run = p.add_run("Agent 面试基础概念救急手册")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run("专门补齐 skills、harness、RAG 召回率、评测指标、工具调用等容易被问懵的基础概念")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    set_run_font(r)


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, color="D7DCE2", size="4")
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    set_run_font(r)
    p.add_run("\n" + body)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_concept_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_borders(table)
    headers = ["概念", "一句话解释", "项目里怎么对应", "面试怎么说"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, LIGHT_BLUE)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_width(table, [1.15, 1.8, 1.65, 1.9])


def add_qa(doc, q, a, bad=None):
    doc.add_paragraph(q, style="Heading 3")
    p = doc.add_paragraph()
    p.add_run("推荐回答：").bold = True
    p.add_run(a)
    if bad:
        p = doc.add_paragraph()
        p.add_run("不要这样答：").bold = True
        p.add_run(bad)


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    add_callout(
        doc,
        "这份文档怎么用",
        "它不是让你变成专家，而是让你面试时听到基础名词不慌。每个概念都按“是什么、为什么重要、项目里怎么对应、怎么口头回答”来准备。遇到不会的细节，不要硬编，可以承认没做深入实验，再说出你知道的判断框架。",
    )

    doc.add_heading("1. 先建立一张 Agent 工程地图", level=1)
    add_bullets(
        doc,
        [
            "用户输入：用户提出任务，例如搜索、分析、写作、计算。",
            "Router：判断任务类型，决定走 ReAct、Plan-Solve、Reflection 或 Supervisor。",
            "Prompt / System Prompt：告诉模型角色、边界、工具使用规则和输出格式。",
            "Tools：模型不能自己访问外部世界，需要通过 search、retrieve、python_repl、MCP 工具等完成动作。",
            "RAG：从本地或企业知识库召回相关资料，再交给模型生成答案。",
            "Memory：保存会话内上下文、长期偏好或摘要，避免每次都像第一次见用户。",
            "Eval / Harness：用一批固定测试题和评分逻辑验证系统是否变好，而不是凭感觉看 demo。",
            "Observability：日志、trace、token、延迟、错误率，用来定位线上问题。",
        ]
    )

    doc.add_heading("2. 高频基础概念速查", level=1)
    add_concept_table(
        doc,
        [
            (
                "Skill",
                "可复用的能力包，通常包含说明、流程、工具或提示词。",
                "你的项目没有严格叫 skill 的模块，但四种 mode 可以理解成不同任务技能；MCP 工具也是外部能力。",
                "我理解 skill 是把某类任务的流程和工具封装起来，让 Agent 遇到相似任务能稳定执行。",
            ),
            (
                "Harness",
                "测试/评测脚手架，用固定流程跑 case 并收集结果。",
                "tests/run_evals.py、benchmarks.py 就是轻量 eval harness。",
                "我把 harness 理解成自动化评测环境，不靠手工看几条输出判断效果。",
            ),
            (
                "RAG",
                "检索增强生成：先查资料，再让模型回答。",
                "docs 文档切块、embedding、Chroma 存储、retrieve 工具召回。",
                "RAG 不是让模型背知识，而是把外部知识动态塞进上下文。",
            ),
            (
                "Recall",
                "检索系统把相关文档找回来的能力。",
                "retrieve(query, k) 从 Chroma 找相关 chunk。",
                "召回率低时，正确答案根本没进上下文，后面模型再强也难答对。",
            ),
            (
                "Precision",
                "召回结果里有多少是真相关。",
                "如果 top-k 里很多无关 chunk，precision 就低。",
                "precision 低会让模型被噪音干扰，容易答偏或幻觉。",
            ),
            (
                "Top-k",
                "每次检索返回前 k 条结果。",
                "retrieve(query, k) 里的 k。",
                "k 小可能漏召回，k 大会带来噪声和 token 成本。",
            ),
            (
                "Chunk",
                "文档切分后的文本片段。",
                "RecursiveCharacterTextSplitter 切 1000 字、overlap 200。",
                "chunk 太小缺上下文，太大检索不准且浪费 token。",
            ),
            (
                "Embedding",
                "把文本转成向量，方便语义相似度检索。",
                "text-embedding-3-small 或 OpenAI-compatible embedding。",
                "embedding 解决的是语义相似，不是关键词完全匹配。",
            ),
            (
                "Vector DB",
                "存储和检索向量的数据库。",
                "Chroma，分别用于 rag_docs 和 agent_memory collection。",
                "向量库负责找相似内容，模型负责基于内容组织答案。",
            ),
            (
                "Structured Output",
                "强制模型输出结构化字段。",
                "SupervisorDecision Pydantic schema。",
                "它能降低解析失败，让 Agent 控制流更稳定。",
            ),
        ],
    )

    doc.add_heading("3. RAG 指标：召回率到底怎么讲", level=1)
    add_callout(
        doc,
        "最重要的一句话",
        "RAG 的质量不是只看最终回答好不好，而是拆成两段：检索有没有把正确材料找回来，生成有没有基于材料答对。召回率主要评估第一段。",
        fill=WARN,
    )
    add_qa(
        doc,
        "Q1：什么是 RAG 召回率？",
        "召回率可以理解为：应该被找到的相关资料，有多少真的被检索系统找回来了。比如一个问题需要命中 10 个相关 chunk，top-k 结果里找回了 7 个，召回率就是 70%。面试里不用死记公式，但要讲清楚它衡量的是“漏没漏掉正确材料”。",
        "召回率就是回答正确率。",
    )
    add_qa(
        doc,
        "Q2：召回率和准确率/精确率有什么区别？",
        "召回率关注“该找的有没有找全”，精确率关注“找回来的里面有多少是有用的”。高召回低精确说明资料找得多但噪声大；高精确低召回说明结果很干净但可能漏掉关键材料。RAG 通常要在两者之间平衡。",
    )
    add_qa(
        doc,
        "Q3：如果 RAG 效果不好，你怎么排查？",
        "我会分层排查：第一，看文档是否解析和切块正确；第二，看 query 是否适合直接检索，是否需要 query rewrite；第三，看 embedding 模型和相似度阈值；第四，看 top-k 是否过小或过大；第五，看最终 prompt 是否把召回内容有效交给模型。",
    )
    add_qa(
        doc,
        "Q4：怎么提升 RAG 召回率？",
        "常见方法有：优化 chunk size 和 overlap，引入 query rewrite，增加 top-k，使用混合检索 BM25 + 向量检索，做 rerank，补元数据过滤，或者针对业务词表做同义词扩展。我的项目目前主要是向量检索，后续可以补混合检索和 rerank。",
    )
    add_qa(
        doc,
        "Q5：你的项目有没有严格测 RAG 召回率？",
        "可以诚实说：我目前做了 benchmark 框架，但 RAG 召回率还没有做成标准化离线评测集。现在更多是验证 retrieve 工具链路和回答效果。如果继续完善，我会构建 query、相关 chunk 标注、top-k hit rate、recall@k、MRR 这些指标。",
    )

    doc.add_heading("4. Harness / Eval：别被这个词吓住", level=1)
    add_callout(
        doc,
        "一句话理解",
        "Harness 就是把“怎么跑测试、跑哪些样例、怎么打分、怎么输出报告”固定下来。它不是某个神秘算法，而是工程化评测流程。",
    )
    add_qa(
        doc,
        "Q1：你的 eval harness 是什么？",
        "我的项目里 tests/benchmarks.py 定义 benchmark case，tests/run_evals.py 负责跑这些 case、调用 graph、收集 mode、tool、关键词和推理深度，再按权重打分。所以它是一个轻量评测 harness。",
    )
    add_qa(
        doc,
        "Q2：为什么 Agent 需要 eval harness？",
        "因为 Agent 的输出不稳定，只看一两次 demo 很容易误判。harness 可以固定测试集，让每次改 prompt、改路由、改工具后都有可比较的结果。它能帮助判断系统是真的变好，还是只是某条样例碰巧成功。",
    )
    add_qa(
        doc,
        "Q3：你现在的评测有什么不足？",
        "不足是 case 数量还少，质量评分主要靠关键词，开放式回答评估不够精细；RAG 召回率、幻觉率、延迟、token 成本也没有完整纳入。后续可以补人工标注集、LLM-as-judge、recall@k、latency 和 cost 指标。",
    )

    doc.add_heading("5. Skills：HR 或面试官问到时怎么接", level=1)
    add_qa(
        doc,
        "Q1：你怎么理解 Agent 里的 skill？",
        "我理解 skill 是一种可复用能力封装。它可以是一套 prompt、工具、流程和约束，也可以是一个面向特定任务的子 Agent。比如“查资料并总结”可以是 research skill，“执行代码计算”可以是 compute skill。重点是让 Agent 不只是聊天，而是按稳定流程完成一类任务。",
    )
    add_qa(
        doc,
        "Q2：你的项目里有 skills 吗？",
        "严格来说项目里没有命名为 skills 的目录，但有类似思想：ReAct、Reflection、Plan-Solve、Supervisor 是不同执行模式；Researcher、Analyst、Executor 是不同角色能力；MCP 工具则是外部能力扩展。如果要进一步工程化，可以把这些模式和工具包装成更标准的 skill registry。",
    )
    add_qa(
        doc,
        "Q3：Skill 和 Tool 有什么区别？",
        "Tool 更像一个具体动作，比如搜索、计算、检索文档；Skill 更像完成一类任务的方法，可能内部会调用多个 tool。比如“写竞品分析报告”是 skill，它可能用 search、retrieve、summary、chart 等多个工具。",
    )

    doc.add_heading("6. Agent 工程基础追问", level=1)
    qa = [
        (
            "Q1：Prompt、Tool、Memory、RAG 分别管什么？",
            "Prompt 负责约束模型怎么想和怎么输出；Tool 负责让模型访问外部能力；Memory 负责保留历史和偏好；RAG 负责从知识库找资料。四者组合起来，才是一个更完整的 Agent 应用。",
        ),
        (
            "Q2：为什么模型需要工具调用？",
            "模型本身只是生成文本，不会真实查网页、跑代码、读数据库。工具调用让模型能把意图转换成外部动作，再把观察结果拿回来继续推理。",
        ),
        (
            "Q3：Function calling 和 MCP 有什么区别？",
            "function calling 更偏模型调用某个函数的协议形式；MCP 更像统一的外部工具服务协议，可以让不同工具 server 被动态发现和接入。简单说，function calling 管“怎么调用函数”，MCP 管“怎么接入一批外部工具”。",
        ),
        (
            "Q4：Agent 为什么容易失控？",
            "因为它会循环推理和调用工具，如果没有最大轮数、工具权限、错误处理和人工确认，就可能重复调用、成本飙升或执行危险操作。所以工程上要加迭代上限、工具白名单、human-in-the-loop 和日志监控。",
        ),
        (
            "Q5：什么是 hallucination？RAG 能完全解决吗？",
            "hallucination 是模型编造不存在的信息。RAG 能缓解，因为它把外部资料放进上下文，但不能完全解决。如果检索材料错了、没召回、prompt 没约束引用，模型仍然可能胡编。",
        ),
        (
            "Q6：为什么要 streaming？",
            "streaming 是流式输出，用户不用等整个 Agent 跑完才看到结果。对长任务来说，它能改善体验，也方便前端展示 token、工具状态和执行进度。",
        ),
    ]
    for q, a in qa:
        add_qa(doc, q, a)

    doc.add_heading("7. 不会答时的兜底模板", level=1)
    add_callout(
        doc,
        "万能但诚实的回答框架",
        "这个点我没有做过完整生产实践，但我理解它大概解决的是 X 问题。我的项目里目前做到 Y，短板是 Z。如果继续做，我会先用 A 指标验证，再尝试 B 方案优化。",
        fill=WARN,
    )
    add_bullets(
        doc,
        [
            "被问 RAG 召回率：我目前没有完整标注集，但知道可以用 recall@k、hit rate、MRR 来评估。",
            "被问 skill：我项目里没有标准 skill registry，但四种 mode 和 specialist 角色体现了能力封装思想。",
            "被问 harness：我的 run_evals.py 是轻量 harness，后续可以接入更完整的评测平台。",
            "被问生产化：我会补权限、日志、metrics、成本统计、人工确认和更安全的代码执行沙箱。",
            "被问不知道的框架名：我可以说没深入用过，但愿意按官方文档和团队现有方案快速补。",
        ]
    )

    doc.add_heading("8. 面试前 30 分钟速背版", level=1)
    add_concept_table(
        doc,
        [
            ("RAG", "先检索资料，再生成答案。", "retrieve + Chroma。", "RAG 质量要分检索和生成两段看。"),
            ("召回率", "该找的资料找回来多少。", "未来可测 recall@k。", "召回低时答案材料没进上下文。"),
            ("精确率", "找回结果里有用的比例。", "top-k 噪声多少。", "精确低会引入干扰。"),
            ("Harness", "自动跑 case 和打分的评测流程。", "run_evals.py。", "不靠肉眼看 demo。"),
            ("Skill", "一类任务能力封装。", "mode / specialist 类似。", "tool 是动作，skill 是流程。"),
            ("MCP", "外部工具接入协议。", "mcp.py + demo server。", "让工具动态扩展，失败可降级。"),
            ("Memory", "保存历史和偏好。", "短期、长期、摘要。", "注意避免 benchmark 污染。"),
            ("Structured output", "模型输出固定结构。", "SupervisorDecision。", "让控制流稳定可解析。"),
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Agent 面试基础概念救急手册 | skills / harness / RAG recall")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    set_run_font(r)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
