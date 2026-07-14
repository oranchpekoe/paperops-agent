from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/agent_engineering_deep_interview_manual.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 32, 44)
MUTED = RGBColor(90, 98, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WARN = "FFF4D6"


def set_run_font(run, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for k, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{k}"))
        if node is None:
            node = OxmlElement(f"w:{k}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def borders(table, color="D7DCE2", size="6"):
    tbl_pr = table._tbl.tblPr
    tb = tbl_pr.first_child_found_in("w:tblBorders")
    if tb is None:
        tb = OxmlElement("w:tblBorders")
        tbl_pr.append(tb)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tb.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tb.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    r = p.add_run("Agent 工程深水区面试手册")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 37, 69)
    set_run_font(r)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run("覆盖记忆机制、上下文管理、Benchmark 评价框架、训练/微调、工具安全、可观测性和生产化追问")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    set_run_font(r)


def callout(doc, title, body, fill=LIGHT_GRAY):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    borders(t, size="4")
    table_width(t, [6.5])
    cell = t.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    rr = p.add_run(title)
    rr.bold = True
    rr.font.color.rgb = DARK_BLUE
    set_run_font(rr)
    p.add_run("\n" + body)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def qa(doc, q, answer, project=None, risk=None):
    doc.add_paragraph(q, style="Heading 3")
    p = doc.add_paragraph()
    p.add_run("参考回答：").bold = True
    p.add_run(answer)
    if project:
        p = doc.add_paragraph()
        p.add_run("结合你的项目：").bold = True
        p.add_run(project)
    if risk:
        p = doc.add_paragraph()
        p.add_run("注意：").bold = True
        p.add_run(risk)


def matrix(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    borders(t)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, LIGHT_BLUE)
        for r in c.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    table_width(t, widths)


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    callout(
        doc,
        "这份文档的目标",
        "让你从“知道几个名词”升级到“能把 Agent 工程链路讲清楚”。面试里最怕不是不会所有东西，而是概念之间没有关系。这份文档会把记忆、上下文、RAG、工具、评测、训练和生产化串成一张图。",
        WARN,
    )

    doc.add_heading("1. Agent 工程全景：先知道每层在干什么", level=1)
    matrix(
        doc,
        ["层次", "解决的问题", "典型组件", "面试重点"],
        [
            ("任务理解层", "判断用户想干什么", "Router、intent classifier、system prompt", "为什么要路由，路由错了怎么办"),
            ("推理编排层", "决定执行流程", "ReAct、Plan-Solve、Reflection、Supervisor", "控制流、循环上限、失败兜底"),
            ("工具执行层", "访问外部世界", "search、retrieve、python_repl、MCP", "权限、安全、超时、错误处理"),
            ("知识层", "补外部知识", "RAG、向量库、rerank、metadata", "召回率、精确率、chunk、混合检索"),
            ("记忆层", "保留历史和偏好", "short-term、long-term、summary memory", "存什么、不存什么、怎么召回、怎么防污染"),
            ("上下文层", "控制 token 预算", "context window、message trimming、compression", "长对话怎么不爆、信息怎么不丢"),
            ("评测层", "判断改动是否有效", "benchmark、eval harness、LLM-as-judge", "指标、case、回归测试"),
            ("运维层", "线上稳定运行", "trace、日志、metrics、成本、延迟", "可观测性和生产化"),
        ],
        [1.1, 1.65, 1.85, 1.9],
    )

    doc.add_heading("2. 记忆机制：Agent 到底怎么“记住”", level=1)
    callout(
        doc,
        "一句话",
        "记忆不是把所有聊天记录无限塞进 prompt，而是决定哪些信息短期保留、哪些长期存储、哪些压缩成摘要、什么时候召回、什么时候删除。",
    )
    qa(
        doc,
        "Q1：短期记忆、长期记忆、摘要记忆有什么区别？",
        "短期记忆是当前会话里的消息历史，解决上下文连贯；长期记忆是跨会话保存的用户偏好、事实或业务状态，解决下次还能想起来；摘要记忆是把过长历史压缩成简短摘要，解决上下文窗口不够。",
        "你的项目用 LangGraph messages / checkpointer 做短期记忆，用 Chroma 的 agent_memory collection 做长期记忆，用 compress_context 做摘要压缩。",
    )
    qa(
        doc,
        "Q2：长期记忆应该存什么？不应该存什么？",
        "应该存稳定、有复用价值的信息，比如用户偏好、长期目标、明确事实、业务配置。不应该随便存一次性闲聊、敏感隐私、模型猜测、benchmark 合成数据和未经确认的信息。否则长期记忆会变成污染源。",
        "你的项目有自动 extract_facts，也有 remember 工具。面试里要强调：自动记忆很方便，但必须有过滤、去重、删除和隐私边界。",
    )
    qa(
        doc,
        "Q3：记忆召回和 RAG 检索有什么区别？",
        "二者技术上都可能用向量检索，但目标不同。RAG 检索的是知识库材料，比如文档、政策、说明书；记忆召回的是和用户或会话有关的历史事实，比如偏好、目标、上次决定。RAG 偏外部知识，memory 偏个体上下文。",
        "你的项目复用了 Chroma，但用不同 collection 隔离 rag_docs 和 agent_memory，这是一个可以讲的工程取舍。",
    )
    qa(
        doc,
        "Q4：长期记忆可能带来什么问题？",
        "主要有五类：错误记忆、过期记忆、隐私风险、召回噪声、评测污染。解决思路包括置信度、时间戳、来源、用户确认、删除机制、敏感信息过滤和 benchmark_mode 隔离。",
        "你的项目已经做了 benchmark 隔离和去重，但还没有完整的隐私权限和用户可编辑记忆界面。",
    )

    doc.add_heading("3. 上下文机制：不是窗口越大越好", level=1)
    qa(
        doc,
        "Q1：什么是 context window？",
        "context window 是模型一次请求能看到的最大 token 范围，包括 system prompt、历史消息、工具结果、RAG 片段和当前问题。超过窗口就必须裁剪、摘要或分步处理。",
        "面试里可以说：上下文是 Agent 的工作内存，管理不好会带来成本高、延迟高、关键信息被挤掉。",
    )
    qa(
        doc,
        "Q2：长对话怎么处理？",
        "常见做法有保留最近 N 轮、把旧消息压缩成摘要、把重要事实写入长期记忆、对工具结果做截断、对 RAG 结果做 rerank 和去重。核心原则是保留决策所需信息，而不是保留所有原文。",
        "你的项目有 keep_last 和 compress_context 思路，适合说成“先保留最近消息，再压缩旧消息”。",
    )
    qa(
        doc,
        "Q3：上下文压缩有什么风险？",
        "摘要可能丢失细节，也可能把模型误解写成事实。解决方法是保留原始引用、摘要只写已确认事实、重要任务保留关键原文、对高风险场景让用户确认。",
    )
    qa(
        doc,
        "Q4：为什么不能把所有工具结果都塞给模型？",
        "一是 token 成本和延迟会变高，二是噪声会干扰模型，三是可能泄露不必要的数据。工具结果应该做筛选、截断、结构化和引用标记。",
    )

    doc.add_heading("4. Benchmark 评价框架：怎么证明 Agent 变好了", level=1)
    callout(
        doc,
        "面试关键",
        "Agent 评测不能只问“这次回答看起来不错吗”，而要拆成路由是否正确、工具是否用对、检索是否召回、答案是否忠实、任务是否完成、成本和延迟是否可接受。",
        WARN,
    )
    matrix(
        doc,
        ["指标", "衡量什么", "常见算法/口径", "你的项目现状"],
        [
            ("Routing accuracy", "路由是否走对模式", "expected_mode vs actual_mode", "已有 benchmark 维度"),
            ("Tool success", "工具是否用对且成功", "expected_tools、forbidden_tools、error rate", "已有工具维度"),
            ("RAG recall", "正确材料是否找回", "recall@k、hit rate、MRR", "可补充，当前不完整"),
            ("Faithfulness", "答案是否忠于材料", "人工评估、LLM-as-judge、引用核验", "可补充"),
            ("Task success", "任务是否完成", "pass/fail、rubric score", "已有关键词近似"),
            ("Cost", "token 和调用成本", "input/output tokens、tool calls", "可补充"),
            ("Latency", "响应耗时", "p50/p95/p99 latency", "可补充"),
            ("Robustness", "换模型/换表达是否稳定", "多模型、多 query 改写", "已有双模型对比基础"),
        ],
        [1.25, 1.55, 1.75, 1.95],
    )
    qa(
        doc,
        "Q1：你的 benchmark 框架怎么设计？",
        "我会说：它定义了一批固定 case，每个 case 标注 expected_mode、expected_tools、forbidden_tools 和 expected_keywords。运行时统一调用 graph，收集实际路由、工具调用和回答，再按权重打分。它的价值是做回归测试，不靠手工看 demo。",
        "你的 tests/benchmarks.py 和 tests/run_evals.py 就是轻量 eval harness。",
    )
    qa(
        doc,
        "Q2：现在 benchmark 最大不足是什么？",
        "case 数量还少，开放式回答靠关键词评分比较粗；RAG recall、faithfulness、延迟、成本、工具错误率还没有完整纳入。如果继续做，我会补标注数据集、recall@k、LLM-as-judge、人工抽检和 trace 指标。",
    )
    qa(
        doc,
        "Q3：LLM-as-judge 靠谱吗？",
        "它适合做辅助评估，但不能完全替代人工。优点是便宜、可批量、能评开放式回答；缺点是 judge model 也会偏、会受 prompt 影响。严肃场景要用人工标注、小样本抽检和一致性校验。",
    )

    doc.add_heading("5. Agent 的训练：实习面试里要讲清楚边界", level=1)
    callout(
        doc,
        "重要纠偏",
        "大多数 Agent 应用工程并不是训练一个大模型，而是基于已有模型做 prompt、RAG、工具、工作流、评测和少量微调。不要把自己说成做了大模型训练。",
        WARN,
    )
    matrix(
        doc,
        ["方式", "做什么", "适合解决", "不适合解决"],
        [
            ("Prompt engineering", "改系统提示和示例", "角色、格式、工具规则", "模型底层能力不足"),
            ("RAG", "外部知识检索", "知识更新、企业文档问答", "复杂行为风格稳定性"),
            ("Fine-tuning", "用样本调整模型输出习惯", "格式、术语、风格、分类", "补实时知识、复杂工具规划"),
            ("Preference / RL", "用偏好信号优化行为", "更符合人类偏好或任务奖励", "小团队快速落地"),
            ("Agent trajectory tuning", "用执行轨迹训练模型", "更会规划和调用工具", "没有高质量轨迹数据时"),
            ("Workflow engineering", "用代码约束流程", "稳定性、权限、安全", "开放式泛化能力"),
        ],
        [1.35, 1.65, 1.75, 1.75],
    )
    qa(
        doc,
        "Q1：你这个项目训练 Agent 了吗？",
        "诚实回答：没有训练底层模型。我做的是 Agent 应用工程，主要通过 LangGraph 工作流、prompt、工具、RAG、记忆和 benchmark 来提升系统表现。训练模型需要高质量数据、算力和评测闭环，不是这个项目重点。",
        risk="不要说“我训练了 Agent”，除非你真的做了模型 fine-tune 或 RL。",
    )
    qa(
        doc,
        "Q2：什么时候需要 fine-tuning？",
        "当问题主要是输出格式、领域术语、风格一致性、分类边界等可从样本中学习的行为时，可以考虑 fine-tuning。如果问题是知识缺失，优先 RAG；如果问题是工具流程不稳定，优先工作流和 structured output。",
    )
    qa(
        doc,
        "Q3：Agent 能不能通过成功/失败轨迹训练？",
        "可以，这是更进阶的方向。比如收集用户任务、模型计划、工具调用、观察结果、最终答案和人工评分，形成 trajectories，再用于监督微调或偏好优化。但数据质量、奖励设计和安全边界都很难，实习项目一般先做好 eval 和日志采集。",
    )

    doc.add_heading("6. 工具调用与安全：Agent 生产化绕不开", level=1)
    qa(
        doc,
        "Q1：工具调用有哪些风险？",
        "风险包括权限越界、误操作、注入攻击、工具返回不可信、超时、重复调用、成本失控。解决方法包括工具白名单、参数校验、超时重试、权限分级、human-in-the-loop、审计日志和危险操作确认。",
        "你的 python_repl 当前是 eval 白名单方案，只适合内部 demo。面试要主动承认公开部署前需要 subprocess、容器或 RestrictedPython。",
    )
    qa(
        doc,
        "Q2：什么是 prompt injection？",
        "用户或文档里夹带恶意指令，让模型忽略系统规则或泄露信息。RAG 场景尤其常见，因为检索文档可能包含“忽略之前指令”之类文本。缓解方式包括系统提示隔离、工具权限控制、检索内容当作不可信数据、输出审查。",
    )
    qa(
        doc,
        "Q3：MCP 工具接入要注意什么？",
        "MCP 让工具扩展方便，但也扩大了攻击面。要关注 server 来源、权限范围、超时、错误隔离、参数 schema、日志审计。外部工具失败不能拖垮主 Agent。",
        "你的项目做了 MCP 懒加载和失败降级，这是一个可讲亮点。",
    )

    doc.add_heading("7. 可观测性、成本和延迟", level=1)
    qa(
        doc,
        "Q1：Agent 为什么特别需要可观测性？",
        "因为 Agent 不是一次普通函数调用，它可能经历路由、LLM 多轮调用、工具调用、RAG 检索、记忆读写。没有 trace，很难知道错误出在路由、检索、工具、prompt 还是模型本身。",
        "你的项目有 trace 文档和 streaming，但结构化日志、metrics、token usage 还可以继续补。",
    )
    qa(
        doc,
        "Q2：线上应该记录哪些指标？",
        "至少包括请求耗时、模型调用次数、工具调用次数、token 消耗、路由模式、错误类型、RAG 命中文档、最终任务成功率。生产里还会看 p50/p95 延迟、失败率、成本、用户反馈。",
    )
    qa(
        doc,
        "Q3：如果 Agent 太慢太贵，怎么优化？",
        "可以减少不必要的路由和反思轮次，限制工具调用上限，缓存检索结果，缩短 prompt，压缩工具输出，降低 top-k，使用更便宜的模型做路由，复杂任务才用强模型。",
    )

    doc.add_heading("8. 生产化差距：你要主动承认，但说出方案", level=1)
    matrix(
        doc,
        ["短板", "为什么重要", "怎么补", "面试表达"],
        [
            ("Human-in-the-loop", "危险操作不能全自动", "interrupt/approval 节点", "已记录，适合下一步做"),
            ("安全沙箱", "代码执行风险高", "subprocess、容器、资源限制", "当前 demo 可用，生产需升级"),
            ("权限系统", "工具和数据要分用户", "RBAC、tenant isolation", "目前单用户项目未做"),
            ("混合检索", "纯向量可能漏关键词", "BM25 + vector + rerank", "RAG 可增强点"),
            ("结构化日志", "线上排障需要", "structlog、trace id、metrics", "当前 trace 不够生产化"),
            ("数据闭环", "优化需要真实反馈", "用户反馈、标注、eval set", "后续补训练/评测基础"),
        ],
        [1.35, 1.65, 1.7, 1.8],
    )

    doc.add_heading("9. 最后速背：被深挖时的回答骨架", level=1)
    bullets(
        doc,
        [
            "记忆：短期管当前会话，长期管跨会话事实，摘要管上下文窗口；风险是污染、过期、隐私。",
            "上下文：窗口有限，要裁剪、摘要、筛选工具结果和 RAG 片段；不是越长越好。",
            "RAG：先看召回，再看生成；召回差就优化 chunk、query rewrite、top-k、混合检索、rerank。",
            "Benchmark：固定 case + 指标 + 自动跑 + 报告；不要只看 demo。",
            "训练：我的项目没有训练模型，做的是 Agent 应用工程；fine-tune 适合格式和风格，RAG 适合知识，workflow 适合稳定流程。",
            "安全：工具调用要限权、校验、超时、审计，危险动作加人工确认。",
            "生产化：日志、trace、token、延迟、成本、权限、human-in-the-loop 都是下一步。",
        ]
    )

    callout(
        doc,
        "万能兜底",
        "这个点我项目里没有完整生产实现，但我知道它在 Agent 工程里主要解决什么问题。当前我做到的是 X，短板是 Y；如果继续做，我会先补指标和日志，再用小规模 benchmark 验证 Z 方案是否真的有效。",
        WARN,
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Agent 工程深水区面试手册 | memory / context / benchmark / training / production")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    set_run_font(r)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
